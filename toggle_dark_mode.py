import sys
import os
import secrets
import openai
from deep_translator import GoogleTranslator
from PyQt6.QtCore import QTimer
from PyQt6.QtWidgets import (
    QApplication, QMainWindow, QTextEdit, QFileDialog, QMessageBox, QPushButton,
    QWidget, QVBoxLayout, QGridLayout
)
from PyQt6.QtGui import QAction, QFont
from PyQt6.QtPrintSupport import QPrinter
from docx import Document
from PyPDF2 import PdfReader
from cryptography.fernet import Fernet
import speech_recognition as sr
import pyttsx3

# OpenAI API Key
OPENAI_API_KEY = "your-api-key-here"
openai.api_key = OPENAI_API_KEY

# Language Mapping
language_map = {
    'Assamese': 'as', 'Bengali': 'bn', 'Gujarati': 'gu', 'Hindi': 'hi', 'Kannada': 'kn', 'Kashmiri': 'ks',
    'Konkani': 'gom', 'Malayalam': 'ml', 'Manipuri': 'mni-Mtei', 'Marathi': 'mr', 'Nepali': 'ne', 'Odia/Oriya': 'or',
    'Punjabi': 'pa', 'Sanskrit': 'sa', 'Sindhi': 'sd', 'Tamil': 'ta', 'Telugu': 'te', 'Urdu': 'ur', 'Bodo': 'brx',
    'Maithili': 'mai', 'Santhali': 'sat', 'Dogri': 'doi', 'English': 'en',
    'Mandarin Chinese': 'zh-CN', 'Spanish': 'es', 'French': 'fr', 'Arabic': 'ar', 'Russian': 'ru',
    'Portuguese': 'pt', 'German': 'de', 'Japanese': 'ja', 'Korean': 'ko'
}

class TextCraftAI(QMainWindow):
    def __init__(self):
        super().__init__()
        self.text_edit = QTextEdit(self)
        self.text_edit.setFont(QFont("Arial", 12))
        self.setCentralWidget(self.text_edit)
        self.setWindowTitle("TextCraft AI - AI-Powered Text Editor")
        self.setGeometry(100, 100, 900, 600)
        self.create_menu()
        self.key = self.generate_key()
        self.auto_save_timer = QTimer(self)
        self.auto_save_timer.timeout.connect(self.auto_save)
        self.auto_save_timer.start(10000)

        self.engine = pyttsx3.init()
        self.recognizer = sr.Recognizer()

        # Load dark mode setting
        self.is_dark_mode = self.load_dark_mode()
        if self.is_dark_mode:
            self.apply_dark_mode()

    def create_menu(self):
        menu_bar = self.menuBar()
        file_menu = menu_bar.addMenu("File")
        file_menu.addAction(self.create_action("Open", self.open_file))
        file_menu.addAction(self.create_action("Save", self.save_file))
        file_menu.addAction(self.create_action("Save as PDF", self.save_as_pdf))
        file_menu.addAction(self.create_action("Save as DOCX", self.save_as_docx))

        ai_menu = menu_bar.addMenu("AI")
        ai_menu.addAction(self.create_action("Generate with AI", self.generate_with_ai))

        translate_menu = menu_bar.addMenu("Translate")
        translate_menu.addAction(self.create_action("Translate Text", self.show_translation_buttons))

        password_menu = menu_bar.addMenu("Password")
        password_menu.addAction(self.create_action("Generate Password", self.generate_password))

        voice_menu = menu_bar.addMenu("Voice")
        voice_menu.addAction(self.create_action("Start Voice Typing", self.start_voice_typing))
        voice_menu.addAction(self.create_action("Speak Text", self.speak_text))

        view_menu = menu_bar.addMenu("View")
        view_menu.addAction(self.create_action("Toggle Dark Mode", self.toggle_dark_mode))

    def create_action(self, name, function):
        action = QAction(name, self)
        action.triggered.connect(function)
        return action

    def auto_save(self):
        with open("autosave.txt", "w", encoding="utf-8") as file:
            file.write(self.text_edit.toPlainText())

    def open_file(self):
        file_name, _ = QFileDialog.getOpenFileName(
            self, "Open File", "", "Text Files (*.txt);;DOCX Files (*.docx);;PDF Files (*.pdf)"
        )
        if file_name:
            try:
                if file_name.endswith(".txt"):
                    with open(file_name, "r", encoding="utf-8") as file:
                        self.text_edit.setText(file.read())
                elif file_name.endswith(".docx"):
                    doc = Document(file_name)
                    self.text_edit.setText("\n".join([para.text for para in doc.paragraphs]))
                elif file_name.endswith(".pdf"):
                    reader = PdfReader(file_name)
                    text = "".join([page.extract_text() or "" for page in reader.pages])
                    self.text_edit.setText(text)
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Failed to open file: {e}")

    def save_file(self):
        file_name, _ = QFileDialog.getSaveFileName(self, "Save File", "", "Text Files (*.txt)")
        if file_name:
            with open(file_name, "w", encoding="utf-8") as file:
                file.write(self.text_edit.toPlainText())

    def save_as_pdf(self):
        file_name, _ = QFileDialog.getSaveFileName(self, "Save as PDF", "", "PDF Files (*.pdf)")
        if file_name:
            printer = QPrinter()
            printer.setOutputFormat(QPrinter.OutputFormat.PdfFormat)
            printer.setOutputFileName(file_name)
            self.text_edit.document().print(printer)
            QMessageBox.information(self, "Success", "File saved as PDF successfully!")

    def save_as_docx(self):
        file_name, _ = QFileDialog.getSaveFileName(self, "Save as DOCX", "", "Word Files (*.docx)")
        if file_name:
            doc = Document()
            doc.add_paragraph(self.text_edit.toPlainText())
            doc.save(file_name)
            QMessageBox.information(self, "Success", "File saved as DOCX successfully!")

    def generate_password(self):
        password = secrets.token_urlsafe(16)
        encrypted_password = self.encrypt_password(password)
        with open("password.txt", "wb") as file:
            file.write(encrypted_password)
        QMessageBox.information(self, "Generated Password", f"Your new password: {password}")

    def generate_with_ai(self):
        prompt = self.text_edit.toPlainText().strip()
        if not prompt:
            QMessageBox.warning(self, "AI Generation", "Please enter some text to generate content.")
            return
        try:
            response = openai.ChatCompletion.create(
                model="gpt-3.5-turbo", messages=[{"role": "user", "content": prompt}]
            )
            self.text_edit.setText(response["choices"][0]["message"]["content"])
        except Exception as e:
            QMessageBox.critical(self, "Error", f"AI Generation Failed: {e}")

    def show_translation_buttons(self):
        self.translation_window = QWidget()
        self.translation_window.setWindowTitle("Translate Text")
        self.translation_window.setGeometry(200, 200, 400, 300)
        
        layout = QGridLayout()
        row, col = 0, 0

        for language in language_map.keys():
            btn = QPushButton(language)
            btn.clicked.connect(lambda checked, lang=language: self.translate_text(lang))
            layout.addWidget(btn, row, col)
            col += 1
            if col > 2:
                col = 0
                row += 1

        self.translation_window.setLayout(layout)
        self.translation_window.show()

    def translate_text(self, language):
        text = self.text_edit.toPlainText().strip()
        if not text:
            QMessageBox.warning(self, "Translation", "Please enter text to translate.")
            return
        try:
            target_language = language_map[language]
            translated_text = GoogleTranslator(target=target_language).translate(text)
            self.text_edit.setText(translated_text)
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Translation Failed: {e}")

    def generate_key(self):
        key_file = "secret.key"
        if os.path.exists(key_file):
            return open(key_file, "rb").read()
        key = Fernet.generate_key()
        with open(key_file, "wb") as file:
            file.write(key)
        return key

    def encrypt_password(self, password):
        return Fernet(self.key).encrypt(password.encode())

    def start_voice_typing(self):
        with sr.Microphone() as source:
            self.recognizer.adjust_for_ambient_noise(source)
            QMessageBox.information(self, "Listening", "Listening for your speech...")
            audio = self.recognizer.listen(source)
        try:
            speech_text = self.recognizer.recognize_google(audio)
            self.text_edit.setText(speech_text)
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Voice typing failed: {e}")

    def speak_text(self):
        text = self.text_edit.toPlainText().strip()
        if text:
            self.engine.say(text)
            self.engine.runAndWait()
        else:
            QMessageBox.warning(self, "No Text", "No text available to speak.")

    # Dark Mode Toggle
    def toggle_dark_mode(self):
        self.is_dark_mode = not self.is_dark_mode
        if self.is_dark_mode:
            self.apply_dark_mode()
        else:
            self.setStyleSheet("")
        self.save_dark_mode(self.is_dark_mode)

    def apply_dark_mode(self):
        self.setStyleSheet("""
            QWidget {
                background-color: #121212;
                color: #ffffff;
            }
            QTextEdit {
                background-color: #1e1e1e;
                color: #ffffff;
                border: 1px solid #444;
            }
            QMenuBar, QMenu, QMenu::item {
                background-color: #1e1e1e;
                color: #ffffff;
            }
            QPushButton {
                background-color: #333;
                color: #fff;
                border-radius: 5px;
                padding: 5px;
            }
            QPushButton:hover {
                background-color: #555;
            }
        """)

    def save_dark_mode(self, enabled):
        with open("dark_mode.txt", "w") as f:
            f.write("on" if enabled else "off")

    def load_dark_mode(self):
        if os.path.exists("dark_mode.txt"):
            with open("dark_mode.txt", "r") as f:
                return f.read().strip() == "on"
        return False

if __name__ == "__main__":
    app = QApplication(sys.argv)
    editor = TextCraftAI()
    editor.show()
    sys.exit(app.exec())
